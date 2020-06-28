# -*- coding: utf-8 -*-

if __name__ == '__main__':
    import re
    import os
    import json
    import docx

    # 要预处理的文件夹路径
    base_path = "english_pro_module/raw_data/"
    base_path_processed = "english_pro_module/raw_data_processed/"
    # 要替换的中英文符号
    rep = {'，': ',', '。': '.', '？': '?', '！': '!', '；': ';', '：': ':', '（': '(', '）': ')',
           '“': '"', '”': '"'}

    # 格式化字符串
    def format_str(rep, str):
        rep = dict((re.escape(k), v) for k, v in rep.items())
        pattern = re.compile("|".join(rep.keys()))
        str_format = pattern.sub(lambda m: rep[re.escape(m.group(0))], str)

        return str_format

    # 处理txt文件
    def process_txt(file):
        file_list = []
        if not os.path.isdir(file):
            f = open(base_path + "/" + file)
            iter_f = iter(f)
            for line in iter_f:
                str_format = format_str(rep, line)
                file_list.append(str_format)

        file_object = open(base_path + "/" + file, 'w')
        for line in file_list:
            file_object.write(line)
            file_object.write('\n')
        file_object.close()

    # 处理docx文件
    def process_docx(file):
        word = docx.Document(base_path + file)
        word_processed = docx.Document()
        for para in word.paragraphs:
            str_format = format_str(rep, para.text)
            word_processed.add_paragraph(str_format)
        word_processed.save(base_path + file.split('.')[0] + '_processed.docx')

    # 处理json文件
    def process_json(file):
        with open(base_path + file) as f:
            pop_data = json.load(f)
            for pop_dict in pop_data:
                for i in range(len(pop_dict['groupQuestions'])):
                    format_article = format_str(rep, pop_dict['groupQuestions'][i]['article'])
                    pop_dict['groupQuestions'][i]['article'] = format_article
                    for j in range(len(pop_dict['groupQuestions'][i]['questions'])):
                        title = format_str(rep, pop_dict['groupQuestions'][i]['questions'][j]['title'])
                        pop_dict['groupQuestions'][i]['questions'][j]['title'] = title
                        for m in range(len(pop_dict['groupQuestions'][i]['questions'][j]['answers'])):
                            answer = format_str(rep, pop_dict['groupQuestions'][i]['questions'][j]['answers'][m])
                            pop_dict['groupQuestions'][i]['questions'][j]['answers'][m] = answer

            print(pop_data)
            with open(base_path + file, 'w') as f:
                json.dump(pop_data, f)

    # 预处理数据
    def preprocessData():
        process_docx('correction.docx')
        process_txt('phrase.txt')
        process_txt('words.txt')
        process_json('listening.json')

    preprocessData()